from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
# 1.创建一个文档对象
document = Document()  # 新建文档对象
# Document("info.docx") 读取现有的word 建立文档对象

# 2.写入内容
document.add_heading('慕课网简介', level=4)
#样式
style=document.styles.add_style('textstyle',WD_STYLE_TYPE.PARAGRAPH)
print(style.style_id)
print(style.name)
style.font.size=Pt(5)
#删除样式
# document.styles['textstyle'].delete()
# 段落
p1 = document.add_paragraph(
    '慕课网是垂直的互联网IT技能免费学习网站。以独家视频教程、在线编程工具、学习计划、问答社区为核心特色。在这里，你可以找到最好的互联网技术牛人，也可以通过免费的在线公开视频课程学习国内领先的互联网IT技术',style='textstyle')
p1.insert_paragraph_before('!!在段落前插入一个新的段落')
format = p1.paragraph_format
# 左右缩进
format.left_indent = Pt(20)
format.right_indent = Pt(20)
format.first_line_indent = Pt(20)  # 首行缩进
# 行间距
format.line_spacing = 1.5
run = p1.add_run('慕课网课程涵盖前端开发、PHP、Html5、Android、iOS、Swift等IT前沿技术语言')
# 字体、字号、文字颜色
run.font.size = Pt(12)
run.font.name = '微软雅黑'
run.font.color.rgb = RGBColor(235, 33, 24)
run1 = p1.add_run('包括基础课程、实用案例、高级分享三大类型，适合不同阶段的学习人群')
# 加粗、下划线、斜体
run1.bold = True
run1.font.underline = True
run1.font.italic = True

# 插入图片
# document.add_picture('logo.jpg')
# document.add_picture('logo.jpg', Pt(20), Pt(30))
# 插入表格
table = document.add_table(rows=1, cols=3,style='Medium List 2')
header_cells = table.rows[0].cells
header_cells[0].text = '月份'
header_cells[1].text = '预期销售额'
header_cells[2].text = '实际销售额'
# 数据
data = (
    ['一月份', 500, 600],
    ['二月份', 900, 600],
    ['三月份', 1000, 600],
)
for item in data:
    rows_cells=table.add_row().cells
    rows_cells[0].text=item[0]
    rows_cells[1].text=str(item[1])
    rows_cells[2].text=str(item[2])
#获取表格
print(len(document.tables[0].rows))#打印总行数
print(len(document.tables[0].columns))#打印总列数
#cell
print(document.tables[0].cell(0,2).text)
# 3.保存文档
document.save('C:/info.docx')
