import xlrd
import random
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

#1.读取excel
data=xlrd.open_workbook('data3.xlsx')
sheet=data.sheet_by_index(0)#获取工作表

class Question:
    pass

def createQuestion():
    questionlist=[]
    for i in range(sheet.nrows):
        if i>1:
            obj=Question()
            obj.subject=sheet.cell(i,1).value #题目
            obj.questiontype = sheet.cell(i, 2).value  # 题型
            obj.option=[]
            obj.option.append(sheet.cell(i, 3).value)  # a
            obj.option.append(sheet.cell(i, 4).value)  # b
            obj.option.append(sheet.cell(i, 5).value)  # c
            obj.option.append(sheet.cell(i, 6).value)  # d
            obj.score = sheet.cell(i, 7).value  # 分值
            questionlist.append(obj)
    random.shuffle(questionlist)#将序列所有的元素随机排序
    return questionlist

#生成word试卷
def createPaper(filename,papername,questionlist):
    document=Document()
    #页眉页脚的信息
    section=document.sections[0]
    header=section.header
    p1=header.paragraphs[0]
    p1.text=papername
    footer=section.footer
    p2=footer.paragraphs[0]
    p2.text="内部试题，禁止泄漏"
    #试卷基本信息
    title=document.add_heading(papername,level=1)
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p3=document.add_paragraph()
    p3.add_run('姓名：______')
    p3.add_run('所属部门：______')
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #试题信息
    for question in questionlist:
        subject=document.add_paragraph(style='List Number')
        run=subject.add_run(question.subject)
        run.bold=True #加粗
        subject.add_run('【%s】分'% str(question.score))
        random.shuffle(question.option) #打乱选项的顺序
        for index,option in enumerate(question.option):
            document.add_paragraph(('ABCD')[index]+str(option))
    document.save(filename)

for i in range(10):
    questionlist=createQuestion()
    createPaper('paper'+str(i+1)+'.docx','2020年第一季度内部考试',questionlist)